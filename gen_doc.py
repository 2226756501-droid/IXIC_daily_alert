from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

title = doc.add_heading("NASDAQ Monitor - Improvement Summary", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph("Plain-language summary of all project improvements.")
doc.add_paragraph("")

doc.add_heading("Part 1: Completed Improvements (10 items)", level=1)

improvements = [
    ("1. Removed Duplicate Function",
     "The event-finalization function was copy-pasted twice.",
     "Deleted the duplicate copy.",
     "Only one place to maintain. Like two identical keys - threw one away."),

    ("2. Fixed 2027 Labor Day Date",
     "Labor Day 2027 is Sept 6 (Monday), but code had Sept 7 (Tuesday).",
     "Changed 7 to 6.",
     "Now correctly knows market open/close. Previously wasted API calls on holidays."),

    ("3. Unified Sensitivity Settings",
     "Two places to set sensitivity, but only one worked.",
     "Removed the dead setting. Only threshold_config.json remains.",
     "One file to change. Like two TV remotes - kept the one with batteries."),

    ("4. Z-score Now Uses 20-Day Window",
     "Used ALL history. 3-year-old data affected today's judgment.",
     "Now only looks at last 20 trading days.",
     "More responsive. Like judging work by last month, not kindergarten."),

    ("5. Added Logging System",
     "Used print() everywhere. Messages disappear when window closes.",
     "Replaced with Python logging. Categorized and saved to file.",
     "Like a flight recorder - review what happened anytime."),

    ("6. Yahoo API Failure Protection",
     "If Yahoo went down, program crashed completely.",
     "Added try/except. Falls back to last cached local data.",
     "Keeps running when Yahoo is down. Like drinking yesterday's milk."),

    ("7. Chart Threshold Follows Sensitivity",
     "Red lines hardcoded at 2. Changing sensitivity changed text but not charts.",
     "Threshold lines now use the same multiplier.",
     "Charts and text always agree. Like thermostat red mark moving with setting."),

    ("8. Expanded Keyword Matching",
     "Only recognized exact words. User variations were ignored.",
     "Added synonym lists for increase and decrease keywords.",
     "Now understands many variations. Like a smarter voice assistant."),

    ("9. Wrote Comprehensive New Tests",
     "Only 3 test files (26 tests). Core modules had NO tests.",
     "Added tests for data_fetcher, visualizer, stats. Total: 47 tests.",
     "Like quality inspection. Changes that break things get caught."),

    ("10. Eliminated Circular Import",
     "analyzer.py and data_fetcher.py risked circular import deadlock.",
     "Extracted shared math into stats.py that depends on nothing.",
     "Clean one-way dependencies. Like clear reporting structure."),
]

for t, problem, fix, effect in improvements:
    doc.add_heading(t, level=2)
    p = doc.add_paragraph()
    run = p.add_run("Problem: ")
    run.bold = True
    p.add_run(problem)
    p = doc.add_paragraph()
    run = p.add_run("Fix: ")
    run.bold = True
    p.add_run(fix)
    p = doc.add_paragraph()
    run = p.add_run("Result: ")
    run.bold = True
    p.add_run(effect)
    doc.add_paragraph("")

doc.add_page_break()
doc.add_heading("Part 2: Remaining Improvements (4 items)", level=1)

remaining = [
    ("1. Docker Multi-stage Build", "Easy", "Low",
     "Image includes docs, tests, git history - ~500MB.",
     "Use multi-stage build, copy only runtime files.",
     "Shrinks to ~300MB, faster download."),

    ("2. Separate CI Test Workflow", "Easy", "Medium",
     "Tests only run in daily workflow.",
     "Create test.yml triggered on every git push.",
     "Instant feedback on every commit."),

    ("3. Local Cache for Dashboard", "Medium", "Medium",
     "Dashboard loads from GitHub raw. Page goes blank if unreachable.",
     "Cache on first load. Use cache when GitHub fails.",
     "Stays visible even with network issues."),

    ("4. Instant Messaging Alerts", "Medium", "Medium",
     "Only email. Can be delayed or go to spam.",
     "Add DingTalk, WeChat Work, or Telegram.",
     "Multiple channels ensure no misses."),
]

for t, diff, imp, cur, sug, eff in remaining:
    doc.add_heading(t, level=2)
    doc.add_paragraph("Difficulty: " + diff)
    doc.add_paragraph("Importance: " + imp)
    p = doc.add_paragraph()
    run = p.add_run("Current: ")
    run.bold = True
    p.add_run(cur)
    p = doc.add_paragraph()
    run = p.add_run("Suggestion: ")
    run.bold = True
    p.add_run(sug)
    p = doc.add_paragraph()
    run = p.add_run("Result: ")
    run.bold = True
    p.add_run(eff)
    doc.add_paragraph("")

doc.save("docs/summary_report.docx")
print("Word doc generated OK")
