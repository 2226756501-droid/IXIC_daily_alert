# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt

doc = Document()

doc.add_heading('NASDAQ 智能监控系统 - 项目详细介绍', 0)

p = doc.add_paragraph()
p.add_run('作者: 田佳祥').bold = True
doc.add_paragraph('（非技术人员，借助 AI 完成）')
doc.add_paragraph('本文档用最白话的语言，把整个项目从头到尾讲清楚。')

# ===== 第一部分 =====
doc.add_heading('一、这是什么项目？', 1)

doc.add_paragraph(
    '这是一个自动监控美国纳斯达克指数（IXIC）的系统。'
    '每天早上它会自动运行，帮你盯着纳斯达克的行情，'
    '如果发现异常（比如连续大跌），它会发邮件通知你。'
    '你还可以随时打开一个网页，看到各种走势图和数据。'
)

doc.add_paragraph('简单说就是：')
doc.add_paragraph('每天早上打开邮箱，就能看到纳斯达克昨天的收盘情况', style='List Bullet')
doc.add_paragraph('如果连续跌了好几天，邮件会告诉你「出事了」，还会附带相关新闻', style='List Bullet')
doc.add_paragraph('随时打开网页，能看到 K 线图、统计分析和 AI 智能问答', style='List Bullet')
doc.add_paragraph('这一切都是自动的，你什么都不用管', style='List Bullet')

doc.add_paragraph('在线体验: https://ixic-daily-alert.streamlit.app')

# ===== 第二部分 =====
doc.add_heading('二、每天会发生什么？', 1)
doc.add_paragraph('这个系统每天晚上（美国时间收盘后）会自动运行一次，做以下几件事：')

doc.add_heading('第1步：抓数据', 2)
doc.add_paragraph(
    '系统会自动访问 Yahoo Finance（雅虎财经），获取纳斯达克指数当天的收盘价、涨跌幅、成交量等数据。'
    '如果雅虎网站连不上（比如网络波动），系统会用之前存好的历史数据继续工作，不会崩溃。'
    '而且如果失败，会自动重试最多 3 次，每次等待时间递增（1秒、2秒、4秒）。'
)

doc.add_heading('第2步：算指标', 2)
doc.add_paragraph(
    '系统会计算一个叫 Z-score 的指标，用来判断今天的涨跌幅在过去 20 天里算不算「异常」。'
    'Z-score 的含义是这样的：'
)

doc.add_paragraph('Z 在 -1 ~ 1 之间：正常波动，不用管', style='List Bullet')
doc.add_paragraph('Z 在 -2 ~ -1 或 1 ~ 2：值得注意，稍微盯着点', style='List Bullet')
doc.add_paragraph('Z 在 -3 ~ -2 或 2 ~ 3：显著异常，要警惕了', style='List Bullet')
doc.add_paragraph('Z 超过 3 或低于 -3：极端行情，很少见', style='List Bullet')

doc.add_paragraph(
    '另外，系统还会看当前的「波动率环境」——如果最近市场波动特别小，那同样一个 Z 值意味着更值得警惕。'
    '还会看成交量是否异常放大（放量下跌比缩量下跌更危险）。'
)

doc.add_heading('第3步：判断是否异常', 2)
doc.add_paragraph('如果纳斯达克连续下跌 3 天，系统会把它标记为「异常时段」，并：')
doc.add_paragraph('在 memory.json 文件里记录这次异常事件（日期、Z值、跌了多少）', style='List Bullet')
doc.add_paragraph('从网上抓取当天的相关新闻，附在邮件里', style='List Bullet')

doc.add_paragraph(
    '如果继续跌到第 4 天，还会计算最近 3 个月的最大回撤（就是从最高点跌了多少百分比），也放在邮件里。'
    '如果行情反弹（涨了），系统会自动结束异常时段，发一封「警报解除」的邮件。'
)

doc.add_heading('第4步：发邮件', 2)
doc.add_paragraph(
    '每天都会发一封邮件到你的 QQ 邮箱。如果配置了 DeepSeek API Key，邮件内容会由 AI 自动生成，'
    '读起来像人写的分析报告，而不是干巴巴的数字。如果 AI 不可用（比如没配置 Key 或网络超时），'
    '系统会自动降级为模板邮件，保证你每天都能收到邮件。'
)

doc.add_paragraph('邮件内容包含：')
doc.add_paragraph('当日收盘价、涨跌幅', style='List Bullet')
doc.add_paragraph('Z-score 和异常等级', style='List Bullet')
doc.add_paragraph('如果处于异常时段，会有醒目的异常提醒', style='List Bullet')
doc.add_paragraph('如果之前发生过类似情况，会告诉你历史上是怎么走的', style='List Bullet')
doc.add_paragraph('邮件末尾有一行反馈：回复 1=满意 2=不满意，帮助改进 AI 生成质量', style='List Bullet')

# ===== 第三部分 =====
doc.add_heading('三、网页上能看到什么？', 1)
doc.add_paragraph('打开网页后，你会看到 6 个标签页：')

doc.add_heading('1. 行情走势', 2)
doc.add_paragraph(
    '这里显示最近 90 天的 K 线图（红色绿色交替的蜡烛图），Z-score 柱状图，以及涨跌幅和 Z-score 的对比图。'
    'K 线图上红色标记的是异常点位。'
)

doc.add_heading('2. 统计分析', 2)
doc.add_paragraph(
    '你可以切换查看 30 天、90 天或全部数据的统计图表，包括收盘价分布、涨跌幅分布、'
    '各月平均收益率、各周涨跌幅分布等。旁边还有具体的数字指标（平均收盘价、标准差等）。'
)

doc.add_heading('3. 回撤分析', 2)
doc.add_paragraph(
    '展示历史上从最高点跌了多少（回撤率），可以看到每次大跌的幅度和持续时间。'
)

doc.add_heading('4. 新闻', 2)
doc.add_paragraph(
    '展示当天跟纳斯达克相关的新闻标题，数据来自雅虎财经和谷歌新闻的 RSS 源。'
)

doc.add_heading('5. 异常事件', 2)
doc.add_paragraph(
    '一个表格，记录了历史上每次异常事件的详细信息：触发日期、Z值、连跌天数、收盘价、持续天数、恢复日期等。'
)

doc.add_heading('6. AI 分析', 2)
doc.add_paragraph(
    '你可以直接打字问 AI 任何关于纳斯达克的问题，比如「今天纳斯达克怎么样？」'
    '「最近一个月表现如何？」「历史上类似情况怎么走的？」'
    'AI 会自动去查实时数据、历史数据、新闻，然后给你一个中文回答。'
)

# ===== 第四部分 =====
doc.add_heading('四、用了哪些技术？', 1)
doc.add_paragraph('下面列出这个项目用到的所有技术，以及它们分别干什么：')

tbl = doc.add_table(rows=10, cols=2)
tbl.style = 'Table Grid'
cells_data = [
    ('技术', '用途'),
    ('Python 3.12', '整个项目用的编程语言，就像盖房子的砖头'),
    ('Streamlit', '把 Python 代码变成网页的工具，不需要会 HTML'),
    ('Yahoo Finance API', '免费的金融数据接口，纳斯达克的数据从这里拿'),
    ('DeepSeek V4 Flash', 'AI 的大脑，理解问题、生成回答，比 ChatGPT 便宜 100 倍'),
    ('OpenAI Agents SDK', '让 AI 能有条理地分析问题、调用工具的框架'),
    ('GitHub Actions', '定时机器人，每天自动运行代码、更新数据'),
    ('Streamlit Cloud', '免费把你的代码变成网页，全世界都能访问'),
    ('Plotly', '画图工具，K线图、柱状图、统计图都是它画的'),
    ('python-dotenv', '管理密钥和配置（API Key、邮箱密码等）'),
]
for i, (c1, c2) in enumerate(cells_data):
    tbl.rows[i].cells[0].text = c1
    tbl.rows[i].cells[1].text = c2

doc.add_heading('这些技术全是免费的：', 2)
doc.add_paragraph('GitHub 仓库：免费', style='List Bullet')
doc.add_paragraph('GitHub Actions：免费（每月 2000 分钟，够用）', style='List Bullet')
doc.add_paragraph('Streamlit Cloud：免费', style='List Bullet')
doc.add_paragraph('Yahoo Finance API：免费', style='List Bullet')
doc.add_paragraph('DeepSeek API：新用户送 5 美元（约 500 万 tokens），每次聊天约 0.2 分钱', style='List Bullet')
doc.add_paragraph('QQ 邮箱：免费（需要开通 SMTP 服务）', style='List Bullet')

# ===== 第五部分 =====
doc.add_heading('五、项目文件都是干什么的？', 1)
doc.add_paragraph('项目根目录下的文件：')

root_files = [
    ('nasdaq.py', '核心脚本，每天自动运行的监控逻辑。抓数据、算指标、判断异常、发邮件'),
    ('app.py', '网页程序，Streamlit 运行的入口。打开这个文件就能启动网页'),
    ('adjust_threshold.py', '通过 GitHub Issue 评论调整报警灵敏度的工具'),
    ('backfill_ohlc.py', '历史数据回填脚本。如果旧数据缺少开盘价、最高价、最低价、成交量，用它补全'),
    ('run_web.bat', 'Windows 下一键启动网页的批处理文件，双击就能打开网页'),
    ('requirements.txt', '项目依赖的 Python 库列表。安装依赖时就靠它'),
    ('Dockerfile', 'Docker 容器构建文件。想用 Docker 部署时用这个'),
    ('threshold_config.json', '灵敏度倍率配置。数值越小越敏感（更容易报警），越大越迟钝'),
    ('history.csv', '每日 NASDAQ 数据记录表。一行是一天的数据'),
    ('market_state.json', '当前市场状态（是否异常时段、连跌几天等）'),
    ('memory.json', '历史异常事件记录'),
    ('.env', '密钥配置（DeepSeek Key、邮箱密码等），不上传到 GitHub'),
]
for name, desc in root_files:
    p = doc.add_paragraph()
    run = p.add_run(name)
    run.bold = True
    p.add_run(' -- ' + desc)

doc.add_paragraph('modules/ 文件夹里的模块文件：')

mod_files = [
    ('data_fetcher.py', '数据抓取，文件读写。负责跟 Yahoo Finance 通信、读写 CSV 和 JSON 文件'),
    ('stats.py', 'Z-score 计算、波动率分析、量能比计算、历史查询'),
    ('analyzer.py', '异常事件记录（进入异常时记录，恢复时结算）'),
    ('news_fetcher.py', '从雅虎和谷歌新闻 RSS 抓取 NASDAQ 相关新闻'),
    ('visualizer.py', '所有图表的绘制（K线图、Z-score 图、统计图、回撤图）'),
    ('drawdown.py', '最大回撤计算（从最高点跌了多少百分比）'),
    ('holidays.py', '美股休市日历，判断今天是不是交易日'),
    ('config.py', '环境变量读取（邮箱配置、API Key 等）'),
    ('agent_engine.py', 'AI Agent 引擎。包含聊天问答和 AI 生成邮件两个功能'),
]
for name, desc in mod_files:
    p = doc.add_paragraph()
    run = p.add_run(name)
    run.bold = True
    p.add_run(' -- ' + desc)

doc.add_paragraph('其他文件：')
other_files = [
    ('tests/', '52 个自动化测试。改了代码后跑一下，能立刻知道有没有改坏'),
    ('.github/workflows/', '3 个 GitHub Actions 自动任务配置（日常运行、测试、灵敏度反馈）'),
    ('docs/', '项目文档'),
]
for name, desc in other_files:
    p = doc.add_paragraph()
    run = p.add_run(name)
    run.bold = True
    p.add_run(' -- ' + desc)

# ===== 第六部分 =====
doc.add_heading('六、数据是怎么流动的？', 1)
doc.add_paragraph('简单说就是一条线：雅虎财经到历史记录到 GitHub 到网页。')

p = doc.add_paragraph()
run = p.add_run(
    'Yahoo Finance -> history.csv -> GitHub Actions (每日运行)\n'
    '                              |-> GitHub Raw (数据源) -> Streamlit Cloud (网页)\n'
    '                              |-> 邮件通知 (每天发送)'
)
run.font.name = 'Courier New'
run.font.size = Pt(10)

doc.add_paragraph(
    '具体过程是：每天 GitHub Actions 启动 -> 运行 nasdaq.py -> '
    '从雅虎财经拿数据 -> 计算 Z-score -> 更新 CSV -> '
    '判断是否异常 -> 发邮件 -> 把 CSV 提交回 GitHub 仓库 -> '
    'Streamlit Cloud 从 GitHub 读取 CSV 显示在网页上。'
)

# ===== 第七部分 =====
doc.add_heading('七、怎么在我自己的电脑上运行？', 1)

doc.add_heading('第1步：下载代码', 2)
p = doc.add_paragraph()
run = p.add_run('git clone https://github.com/2226756501-droid/IXIC_daily_alert.git')
run.font.name = 'Courier New'
run.font.size = Pt(10)

doc.add_heading('第2步：安装依赖', 2)
for cmd in [
    'cd IXIC_daily_alert',
    'python -m venv .venv        # 创建虚拟环境（可选但推荐）',
    '.venv\\Scripts\\activate     # 激活虚拟环境（Windows）',
    'pip install -r requirements.txt',
]:
    p = doc.add_paragraph()
    run = p.add_run(cmd)
    run.font.name = 'Courier New'
    run.font.size = Pt(10)

doc.add_heading('第3步：配置 .env 文件', 2)
doc.add_paragraph('复制 .env.example 为 .env，填上以下信息（不填也能运行，只是 AI 和邮件功能不可用）：')
doc.add_paragraph('DEEPSEEK_API_KEY：你的 DeepSeek API Key（用于 AI 聊天和生成邮件）', style='List Bullet')
doc.add_paragraph('EMAIL_USER / EMAIL_PASS：QQ 邮箱账号和 SMTP 授权码（用于发邮件）', style='List Bullet')
doc.add_paragraph('NOTIFY_EMAIL：接收通知的邮箱地址', style='List Bullet')

doc.add_heading('第4步：启动网页', 2)
p = doc.add_paragraph()
run = p.add_run('streamlit run app.py')
run.font.name = 'Courier New'
run.font.size = Pt(10)
doc.add_paragraph('浏览器打开 http://localhost:8501 就能看到网页。')

doc.add_heading('第5步（可选）：运行测试', 2)
p = doc.add_paragraph()
run = p.add_run('python -m pytest tests/ -v')
run.font.name = 'Courier New'
run.font.size = Pt(10)
doc.add_paragraph('会显示 52 个测试全部通过。')

doc.add_heading('第6步（可选）：回填历史数据', 2)
doc.add_paragraph(
    '如果 history.csv 里的旧数据缺少开盘价、最高价、最低价、成交量（K 线图会回退为折线图），'
    '可以运行回填脚本补全：'
)
p = doc.add_paragraph()
run = p.add_run('python backfill_ohlc.py')
run.font.name = 'Courier New'
run.font.size = Pt(10)

# ===== 第八部分 =====
doc.add_heading('八、常见问题', 1)

qa_pairs = [
    ('问：这个项目要花钱吗？',
     '几乎不花。GitHub 和 Streamlit Cloud 都是免费的。DeepSeek API 新用户送 5 美元，'
     '每次聊天大约 0.2 分钱，一年可能用不到 1 块钱。QQ 邮箱免费。唯一可能需要花钱的是你的服务器（如果你不用免费的 Streamlit Cloud）。'),
    ('问：我不会编程，能自己部署吗？',
     '可以。你只需要按上面的第1步到第4步在电脑上执行命令即可，不需要写代码。'
     '但如果想部署到 Streamlit Cloud 和配置 GitHub Actions，需要一点 Git 基础。'),
    ('问：邮件发不出来怎么办？',
     '检查 .env 里的邮箱配置是否正确。QQ 邮箱需要开启 SMTP 服务，获取授权码（不是登录密码）。'
     '如果没配邮箱，系统会跳过发邮件，不影响其他功能。'),
    ('问：AI 功能用不了？',
     '需要配置 DEEPSEEK_API_KEY。如果没配，AI 聊天会提示「功能未配置」，邮件会使用模板生成（不依赖 AI）。'),
    ('问：数据准不准？',
     '数据来自雅虎财经（Yahoo Finance），这是全球最大的金融数据提供商之一，数据是可靠的。'
     '但如果雅虎偶尔出问题，系统会用缓存数据兜底，保证不崩溃。'),
    ('问：我能修改报警的灵敏度吗？',
     '可以。在 threshold_config.json 里修改 sensitivity_multiplier 的值。'
     '默认是 1.0，调小（比如 0.5）会更容易触发报警，调大（比如 2.0）会更迟钝。'
     '也可以通过 GitHub Issue 评论来调整（通过 feedback.yml 工作流）。'),
    ('问：这个项目是谁做的？',
     '作者是田佳祥，一个不会编程的非技术人员。'
     '整个项目包括代码、文档、改进建议，全部是通过跟 AI 对话完成的。'
     '每次改进的思路都是：问 AI 还有什么可以改进 -> AI 给建议 -> 点头或摇头 -> 实现。'
     '这个项目证明了：不懂代码的人也能做出一个完整的上线项目。'),
]
for q, a in qa_pairs:
    doc.add_heading(q, 2)
    doc.add_paragraph(a)

# ===== 第九部分 =====
doc.add_heading('九、版本历史', 1)

versions = [
    ('最初版', '每天抓个数字发邮件'),
    ('V2', '加了网页图表，能看走势图'),
    ('V3', '加了 Z-score 异常检测，自动标记异常时段'),
    ('V4', '加了灵敏度调节（threshold_config.json），本地缓存兜底'),
    ('V5', '加了 Docker 支持，最大回撤分析'),
    ('V6', '加了 AI 聊天（网页上直接问 AI 问题）'),
    ('V7', 'AI 生成邮件内容，替换硬编码拼接'),
    ('V8', 'K 线图替换折线图，统计分析分时间范围'),
    ('V9', '波动率聚类 + 成交量确认，更精准的异常检测'),
    ('V10（最新）', '邮件反馈闭环，指数退避重试，网页缓存，历史数据回填'),
]
for ver, desc in versions:
    p = doc.add_paragraph()
    run = p.add_run(ver)
    run.bold = True
    p.add_run(' -- ' + desc)

# ===== 保存 =====
output_path = 'D:\\ixic\\docs\\IXIC项目详细介绍_小白版.docx'
doc.save(output_path)
print(f'文档已生成: {output_path}')
